"""Document processor for chunking and preprocessing."""

import re
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from enum import Enum

from services.common.logging import get_logger

logger = get_logger(__name__)


class ChunkingStrategy(Enum):
    """Chunking strategies."""
    FIXED_SIZE = "fixed_size"
    SENTENCE = "sentence"
    PARAGRAPH = "paragraph"
    SEMANTIC = "semantic"
    RECURSIVE = "recursive"


@dataclass
class DocumentChunk:
    """Document chunk."""
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0


@dataclass
class ProcessedDocument:
    """Processed document with chunks."""
    original_content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentProcessor:
    """
    Process documents for indexing.

    Handles:
    - Text cleaning and normalization
    - Document chunking
    - Metadata extraction
    """

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE,
        min_chunk_size: int = 100,
        separators: Optional[List[str]] = None,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            chunking_strategy: Chunking strategy to use
            min_chunk_size: Minimum chunk size
            separators: Custom separators for recursive chunking
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.min_chunk_size = min_chunk_size
        self.separators = separators or ["\n\n", "\n", ". ", " ", ""]

    def process(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """
        Process a document.

        Args:
            content: Document content
            metadata: Document metadata

        Returns:
            ProcessedDocument with chunks
        """
        # Clean content
        cleaned_content = self._clean_text(content)

        # Extract additional metadata
        extracted_metadata = self._extract_metadata(cleaned_content)
        if metadata:
            extracted_metadata.update(metadata)

        # Chunk document
        chunks = self._chunk_document(cleaned_content)

        # Add metadata to chunks
        for chunk in chunks:
            chunk.metadata.update(extracted_metadata)

        return ProcessedDocument(
            original_content=content,
            chunks=chunks,
            metadata=extracted_metadata,
        )

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove null bytes
        text = text.replace('\x00', '')

        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        """Extract metadata from text."""
        metadata = {}

        # Word count
        metadata["word_count"] = len(text.split())

        # Character count
        metadata["char_count"] = len(text)

        # Language detection (simple heuristic)
        metadata["language"] = self._detect_language(text)

        return metadata

    def _detect_language(self, text: str) -> str:
        """Simple language detection."""
        # Check for Chinese characters
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        total_chars = len(text)

        if total_chars > 0 and chinese_chars / total_chars > 0.3:
            return "zh"
        return "en"

    def _chunk_document(self, text: str) -> List[DocumentChunk]:
        """Chunk document based on strategy."""
        if self.chunking_strategy == ChunkingStrategy.FIXED_SIZE:
            return self._fixed_size_chunking(text)
        elif self.chunking_strategy == ChunkingStrategy.SENTENCE:
            return self._sentence_chunking(text)
        elif self.chunking_strategy == ChunkingStrategy.PARAGRAPH:
            return self._paragraph_chunking(text)
        elif self.chunking_strategy == ChunkingStrategy.RECURSIVE:
            return self._recursive_chunking(text)
        else:
            return self._fixed_size_chunking(text)

    def _fixed_size_chunking(self, text: str) -> List[DocumentChunk]:
        """Fixed size chunking with overlap."""
        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + self.chunk_size

            # Find a good break point
            if end < len(text):
                # Try to break at space
                break_point = text.rfind(' ', start, end)
                if break_point > start + self.min_chunk_size:
                    end = break_point

            chunk_text = text[start:end].strip()

            if len(chunk_text) >= self.min_chunk_size:
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                ))
                chunk_index += 1

            start = end - self.chunk_overlap

        return chunks

    def _sentence_chunking(self, text: str) -> List[DocumentChunk]:
        """Sentence-based chunking."""
        # Split by sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)

        chunks = []
        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                    ))
                    chunk_index += 1
                    current_start += len(current_chunk)
                current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        if current_chunk and len(current_chunk) >= self.min_chunk_size:
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            ))

        return chunks

    def _paragraph_chunking(self, text: str) -> List[DocumentChunk]:
        """Paragraph-based chunking."""
        paragraphs = text.split('\n\n')

        chunks = []
        current_chunk = ""
        current_start = 0
        chunk_index = 0

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            if len(current_chunk) + len(paragraph) > self.chunk_size:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                    ))
                    chunk_index += 1
                    current_start += len(current_chunk)
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph

        if current_chunk and len(current_chunk) >= self.min_chunk_size:
            chunks.append(DocumentChunk(
                content=current_chunk.strip(),
                chunk_index=chunk_index,
                start_char=current_start,
                end_char=current_start + len(current_chunk),
            ))

        return chunks

    def _recursive_chunking(self, text: str) -> List[DocumentChunk]:
        """
        Recursive character text splitting.

        Tries to split on larger semantic boundaries first,
        then falls back to smaller ones.
        """
        return self._recursive_split(text, self.separators, 0)

    def _recursive_split(
        self,
        text: str,
        separators: List[str],
        start_char: int,
    ) -> List[DocumentChunk]:
        """Recursive split helper."""
        chunks = []

        if not separators:
            # Base case: no more separators, use fixed size
            return self._fixed_size_chunking(text)

        separator = separators[0]
        remaining_separators = separators[1:]

        if separator:
            splits = text.split(separator)
        else:
            # Empty separator means character-by-character
            splits = list(text)

        current_chunk = ""
        current_start = start_char
        chunk_index = len(chunks)

        for i, split in enumerate(splits):
            # Add separator back (except for last split)
            piece = split if i == len(splits) - 1 else split + separator

            if len(current_chunk) + len(piece) <= self.chunk_size:
                current_chunk += piece
            else:
                if current_chunk:
                    if len(current_chunk) > self.chunk_size:
                        # Chunk is too big, recurse with smaller separator
                        sub_chunks = self._recursive_split(
                            current_chunk,
                            remaining_separators,
                            current_start,
                        )
                        for sub_chunk in sub_chunks:
                            sub_chunk.chunk_index = chunk_index
                            chunk_index += 1
                        chunks.extend(sub_chunks)
                    elif len(current_chunk) >= self.min_chunk_size:
                        chunks.append(DocumentChunk(
                            content=current_chunk.strip(),
                            chunk_index=chunk_index,
                            start_char=current_start,
                            end_char=current_start + len(current_chunk),
                        ))
                        chunk_index += 1

                    current_start += len(current_chunk)

                current_chunk = piece

        # Handle remaining chunk
        if current_chunk:
            if len(current_chunk) > self.chunk_size:
                sub_chunks = self._recursive_split(
                    current_chunk,
                    remaining_separators,
                    current_start,
                )
                for sub_chunk in sub_chunks:
                    sub_chunk.chunk_index = chunk_index
                    chunk_index += 1
                chunks.extend(sub_chunks)
            elif len(current_chunk) >= self.min_chunk_size:
                chunks.append(DocumentChunk(
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                ))

        # Re-number chunk indices
        for i, chunk in enumerate(chunks):
            chunk.chunk_index = i

        return chunks


class FileProcessor:
    """
    Process different file types.

    Supports:
    - Plain text (.txt)
    - Markdown (.md)
    - PDF (.pdf)
    - Word documents (.docx)
    """

    def __init__(self, document_processor: Optional[DocumentProcessor] = None):
        """
        Initialize file processor.

        Args:
            document_processor: Document processor for chunking
        """
        self.document_processor = document_processor or DocumentProcessor()

    async def process_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """
        Process a file.

        Args:
            file_path: Path to file
            metadata: Additional metadata

        Returns:
            ProcessedDocument
        """
        import os

        extension = os.path.splitext(file_path)[1].lower()

        # Read file content based on type
        if extension in ['.txt', '.md']:
            content = await self._read_text_file(file_path)
        elif extension == '.pdf':
            content = await self._read_pdf_file(file_path)
        elif extension == '.docx':
            content = await self._read_docx_file(file_path)
        else:
            # Try to read as text
            content = await self._read_text_file(file_path)

        # Add file metadata
        file_metadata = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_extension": extension,
        }
        if metadata:
            file_metadata.update(metadata)

        return self.document_processor.process(content, file_metadata)

    async def _read_text_file(self, file_path: str) -> str:
        """Read text file."""
        import aiofiles

        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()

    async def _read_pdf_file(self, file_path: str) -> str:
        """Read PDF file."""
        try:
            import pypdf

            text_parts = []
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return '\n\n'.join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, cannot read PDF files")
            return ""

    async def _read_docx_file(self, file_path: str) -> str:
        """Read Word document."""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n\n'.join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed, cannot read DOCX files")
            return ""
